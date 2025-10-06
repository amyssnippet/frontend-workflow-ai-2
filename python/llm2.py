import subprocess
import os
import requests
import json
import re
import mimetypes
import base64
import io
from typing import List, Dict, Tuple
import asyncio
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
from fastapi import FastAPI, UploadFile, File, Form, HTTPException,Body
from fastapi.responses import JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image, ImageDraw, ImageFont

import pdfplumber
from docx import Document as DocxDocument
from pptx import Presentation

app = FastAPI(title="Ollama Mermaid Diagram API")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_methods=["*"],
    allow_headers=["*"],
)

if not os.path.exists("static"):
    os.makedirs("static")
if not os.path.exists("static/process_images"):
    os.makedirs("static/process_images")
if not os.path.exists("input/images"):
    os.makedirs("input/images")

app.mount("/static", StaticFiles(directory="static"), name="static")

OLLAMA_API_URL = "https://suggested-accuracy-brokers-gossip.trycloudflare.com/api/chat"
OLLAMA_MODEL = "granite3.3:8b"
VISION_MODEL = "qwen2.5vl:3b"
token = "ac01d6d03fe9feac354add11a4ac67ef51a089594835d56056436f48c9e00ed1"

# Concurrency controls: limit to 10 concurrent calculation requests.
# We'll create executors on startup to avoid ProcessPool issues on Windows at import time.
MAX_CONCURRENT_REQUESTS = 10
IO_WORKERS = 20
CPU_WORKERS = 10


@app.on_event("startup")
async def startup_event():
    # Thread pool for I/O bound tasks (network calls)
    app.state.io_executor = ThreadPoolExecutor(max_workers=IO_WORKERS)
    # Process pool for CPU-bound tasks (image/docx generation)
    app.state.cpu_executor = ProcessPoolExecutor(max_workers=CPU_WORKERS)
    # Semaphore to limit number of concurrent requests doing heavy work
    app.state.semaphore = asyncio.Semaphore(MAX_CONCURRENT_REQUESTS)


@app.on_event("shutdown")
async def shutdown_event():
    try:
        app.state.io_executor.shutdown(wait=False)
    except Exception:
        pass
    try:
        app.state.cpu_executor.shutdown(wait=False)
    except Exception:
        pass

def check_mmdc_installed():
    try:
        if os.name == "nt":
            mmdc_path = r"C:\Users\amoly\AppData\Roaming\npm\mmdc.cmd"
        else:
            mmdc_path = r"/usr/bin/mmdc"
        subprocess.run([mmdc_path, "--version"], check=True, capture_output=True)
        return True
    except (FileNotFoundError, subprocess.CalledProcessError):
        print("Error: Mermaid CLI (mmdc) not found.")
        return False

def summarize_process_details(analysis_text: str) -> Dict[str, str]:
    """Extract concise summary from detailed analysis."""
    print(f"[SUMMARIZE] Processing analysis text of length {len(analysis_text)}")
    
    lines = [line.strip() for line in analysis_text.split('\n') if line.strip()]
    
    summary = {
        'action': '',
        'input': '',
        'output': '',
        'key_details': []
    }
    
    # Extract key information from numbered points
    for line in lines:
        clean_line = re.sub(r'^\d+\.\s*', '', line)
        lower_line = clean_line.lower()
        
        # Identify action/process
        if any(keyword in lower_line for keyword in ['action', 'process', 'shown', 'displays', 'interface']):
            if not summary['action']:
                summary['action'] = clean_line[:150]
        
        # Identify inputs/prerequisites
        elif any(keyword in lower_line for keyword in ['input', 'prerequisite', 'required', 'need']):
            if not summary['input']:
                summary['input'] = clean_line[:150]
        
        # Identify outputs/results
        elif any(keyword in lower_line for keyword in ['output', 'result', 'generate', 'create']):
            if not summary['output']:
                summary['output'] = clean_line[:150]
        
        # Collect other important details
        elif any(keyword in lower_line for keyword in ['setting', 'configuration', 'option', 'field', 'button', 'value']):
            if len(summary['key_details']) < 3:
                summary['key_details'].append(clean_line[:100])
    
    # Fallback: use first few lines if specific sections not found
    if not summary['action'] and lines:
        summary['action'] = lines[0][:150]
    
    if not summary['key_details']:
        summary['key_details'] = [re.sub(r'^\d+\.\s*', '', line)[:100] for line in lines[1:4]]
    
    print(f"[SUMMARIZE] Extracted - Action: {summary['action'][:50]}...")
    print(f"[SUMMARIZE] Key details count: {len(summary['key_details'])}")
    
    return summary

def extract_steps_from_analysis(image_analyses: List[Dict]) -> List[Dict]:
    """Extract sequential workflow steps from image analyses."""
    print(f"\n[EXTRACT_STEPS] Extracting steps from {len(image_analyses)} image analyses")
    
    steps = []
    
    for idx, analysis_data in enumerate(image_analyses, 1):
        analysis_text = analysis_data['analysis']
        image_path = analysis_data['image_path']
        filename = analysis_data['filename']
        
        # Summarize the analysis
        summary = summarize_process_details(analysis_text)
        
        # Use action as the main description
        step_description = summary['action'] or f"Process from {filename}"
        
        steps.append({
            "number": idx,
            "title": f"Step {idx}",
            "description": step_description,
            "summary": summary,
            "full_analysis": analysis_text,
            "image_path": image_path,
            "filename": filename,
            "next_step": idx + 1 if idx < len(image_analyses) else None
        })
        
        print(f"[EXTRACT_STEPS] Step {idx}: {step_description[:100]}")
    
    print(f"[EXTRACT_STEPS] Total steps extracted: {len(steps)}")
    return steps

def create_sop_docx_with_images(workflow_steps: List[Dict], flowchart_path: str = None) -> str:
    """Create a clean SOP/Workflow DOCX document with images and step-by-step instructions."""
    
    print(f"\n[CREATE_DOCX] Starting DOCX creation with images")
    print(f"[CREATE_DOCX] Workflow steps: {len(workflow_steps)}")
    
    doc = DocxDocument()
    
    # Document Title
    title = doc.add_heading('Standard Operating Procedure', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Step-by-Step Workflow Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(16)
    subtitle_format.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    print(f"[CREATE_DOCX] Adding {len(workflow_steps)} steps with images to document")
    
    # Add each step with its image
    for step in workflow_steps:
        # Add the screenshot/image for this step FIRST
        if step.get('image_path') and os.path.exists(step['image_path']):
            try:
                # Center align the image
                image_para = doc.add_paragraph()
                image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add image
                run = image_para.add_run()
                run.add_picture(step['image_path'], width=Inches(6.5))
                
                doc.add_paragraph()  # Spacing after image
                
                print(f"[CREATE_DOCX] Added image for step {step['number']}: {step['image_path']}")
            except Exception as e:
                doc.add_paragraph(f"⚠ Unable to add image: {str(e)}")
                print(f"[CREATE_DOCX] Error adding image for step {step['number']}: {e}")
        
        # Add Figure caption
        figure_para = doc.add_paragraph()
        figure_run = figure_para.add_run(f"Figure: {step['description']}")
        figure_run.bold = True
        figure_run.font.size = Pt(11)
        
        doc.add_paragraph()  # Spacing
        
        # Step-by-Step Instructions heading
        if step['number'] == 1:
            instructions_heading = doc.add_paragraph()
            instructions_run = instructions_heading.add_run('Step-by-Step Instructions:')
            instructions_run.bold = True
            instructions_run.font.size = Pt(12)
        else:
            instructions_heading = doc.add_paragraph()
            instructions_run = instructions_heading.add_run('Step-by-Step Instructions (Continued):')
            instructions_run.bold = True
            instructions_run.font.size = Pt(12)
        
        doc.add_paragraph()  # Spacing
        
        # Parse the analysis to extract sub-steps
        analysis_lines = [line.strip() for line in step['full_analysis'].split('\n') if line.strip()]
        
        substep_counter = 1
        for line in analysis_lines:
            # Remove leading numbers and hyphens
            clean_line = re.sub(r'^\d+\.\s*', '', line)
            clean_line = re.sub(r'^-\s*', '', clean_line).strip()
            
            if not clean_line or len(clean_line) < 15:
                continue
            
            # Split by ' - ' if it exists
            if ' - ' in clean_line:
                parts = clean_line.split(' - ', 1)
                action = parts[0].strip()
                details = parts[1].strip()
            else:
                action = clean_line
                details = None
            
            # Add substep with bold step number and action
            substep_para = doc.add_paragraph()
            substep_run = substep_para.add_run(f'Step {substep_counter}: ')
            substep_run.bold = True
            substep_run.font.size = Pt(11)
            substep_para.add_run(action)
            
            # ONLY add bullet if there are additional details different from action
            if details and details.lower() != action.lower():
                bullet_para = doc.add_paragraph(style='List Bullet')
                bullet_para.add_run(f'- {details}')
                bullet_para.paragraph_format.left_indent = Inches(0.5)
            
            substep_counter += 1
        
        doc.add_paragraph()  # Extra spacing between main steps
        
        # Page break after each main step
        if step['number'] < len(workflow_steps):
            doc.add_page_break()
        
        print(f"[CREATE_DOCX] Completed step {step['number']}: {step['description']}")
    
    # Save document
    docx_filename = "workflow_sop.docx"
    docx_path = f"static/{docx_filename}"
    doc.save(docx_path)
    
    print(f"[CREATE_DOCX] DOCX saved to: {docx_path}")
    print(f"[CREATE_DOCX] File size: {os.path.getsize(docx_path)} bytes")
    
    return docx_path

def call_ollama_vision(image_data: bytes, prompt: str) -> str:
    """Call vision model to analyze image."""
    print(f"\n[VISION] Calling vision model with image size: {len(image_data)} bytes")
    print(f"[VISION] Prompt: {prompt[:150]}...")
    
    try:
        encoded_image = base64.b64encode(image_data).decode('utf-8')
        print(f"[VISION] Encoded image size: {len(encoded_image)} characters")
        
        messages = [
            {
                "role": "user",
                "content": prompt,
                "images": [encoded_image]
            }
        ]
        
        payload = {
            "model": VISION_MODEL,
            "messages": messages,
            "stream": False,
            "options": {
                "temperature": 0.1,
                "num_ctx": 4096
            }
        }
        
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {token}"
        }
        
        print(f"[VISION] Sending request to: {OLLAMA_API_URL}")
        response = requests.post(
            OLLAMA_API_URL,
            headers=headers,
            data=json.dumps(payload),
            timeout=120
        )
        
        print(f"[VISION] Response status code: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            content = result.get("message", {}).get("content", "").strip()
            print(f"[VISION] Response length: {len(content)} characters")
            return content
        else:
            print(f"[VISION] Error response: {response.text[:500]}")
            return ""
            
    except Exception as e:
        print(f"[VISION] Exception: {e}")
        import traceback
        traceback.print_exc()
        return ""

def create_process_image(step_number: int, title: str, description: str, status: str, data: Dict = None) -> str:
    """Create a visual representation of a process step."""
    
    print(f"[PROCESS_IMAGE] Creating image for step {step_number}: {title}")
    
    width = 800
    height = 400
    
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    
    try:
        font_title = ImageFont.truetype("arial.ttf", 32)
        font_desc = ImageFont.truetype("arial.ttf", 20)
        font_number = ImageFont.truetype("arialbd.ttf", 80)
        font_status = ImageFont.truetype("arial.ttf", 18)
    except:
        font_title = ImageFont.load_default()
        font_desc = ImageFont.load_default()
        font_number = ImageFont.load_default()
        font_status = ImageFont.load_default()
    
    if status == "success":
        bg_color = (34, 139, 34)
        status_text = "✓ SUCCESS"
    elif status == "processing":
        bg_color = (30, 144, 255)
        status_text = "⟳ PROCESSING"
    elif status == "error":
        bg_color = (220, 20, 60)
        status_text = "✗ ERROR"
    else:
        bg_color = (128, 128, 128)
        status_text = "○ PENDING"
    
    draw.rectangle([(0, 0), (width, 100)], fill=bg_color)
    
    circle_center = (80, 50)
    circle_radius = 40
    draw.ellipse(
        [(circle_center[0] - circle_radius, circle_center[1] - circle_radius),
         (circle_center[0] + circle_radius, circle_center[1] + circle_radius)],
        fill='white'
    )
    
    number_text = str(step_number)
    number_bbox = draw.textbbox((0, 0), number_text, font=font_number)
    number_width = number_bbox[2] - number_bbox[0]
    number_height = number_bbox[3] - number_bbox[1]
    draw.text(
        (circle_center[0] - number_width // 2, circle_center[1] - number_height // 2 - 5),
        number_text,
        fill=bg_color,
        font=font_number
    )
    
    draw.text((140, 30), title, fill='white', font=font_title)
    
    status_bbox = draw.textbbox((0, 0), status_text, font=font_status)
    status_width = status_bbox[2] - status_bbox[0]
    draw.text((width - status_width - 20, 35), status_text, fill='white', font=font_status)
    
    draw.line([(20, 110), (width - 20, 110)], fill=bg_color, width=3)
    
    y_offset = 140
    words = description.split()
    lines = []
    current_line = []
    
    for word in words:
        current_line.append(word)
        test_line = ' '.join(current_line)
        bbox = draw.textbbox((0, 0), test_line, font=font_desc)
        if bbox[2] - bbox[0] > width - 80:
            if len(current_line) > 1:
                current_line.pop()
                lines.append(' '.join(current_line))
                current_line = [word]
            else:
                lines.append(test_line)
                current_line = []
    
    if current_line:
        lines.append(' '.join(current_line))
    
    for line in lines[:5]:
        draw.text((40, y_offset), line, fill='black', font=font_desc)
        y_offset += 30
    
    if data:
        y_offset += 20
        draw.line([(40, y_offset), (width - 40, y_offset)], fill='lightgray', width=1)
        y_offset += 15
        
        for key, value in list(data.items())[:3]:
            text = f"{key}: {value}"
            if len(text) > 80:
                text = text[:77] + "..."
            draw.text((40, y_offset), text, fill='gray', font=font_status)
            y_offset += 25
    
    draw.rectangle([(0, 0), (width - 1, height - 1)], outline=bg_color, width=3)
    
    filename = f"process_step_{step_number:02d}.png"
    filepath = f"static/process_images/{filename}"
    img.save(filepath, 'PNG')
    
    print(f"[PROCESS_IMAGE] Saved to {filepath}")
    return filename

@app.post("/generate-process-images/")
async def generate_diagram_with_process_images(
    description: str = Form(None),
    files: List[UploadFile] = File(None),
    output_format: str = Form("png"),
):
    """Generate SOP from multiple images treating each as a sequential step."""
    
    print("\n" + "="*80)
    print("STARTING NEW REQUEST - SEQUENTIAL IMAGE PROCESSING")
    print("="*80)
    
    # Acquire semaphore so only MAX_CONCURRENT_REQUESTS are processed in parallel
    async with app.state.semaphore:
        # Clear previous process images
        process_dir = "static/process_images"
        for f in os.listdir(process_dir):
            if f.startswith("process_step_"):
                os.remove(os.path.join(process_dir, f))

        process_steps = []
        image_analyses = []
        step_counter = 1

        # Process Multiple Images as Sequential Steps
        if files and len(files) > 0:
            print(f"\n[MAIN] Processing {len(files)} images as sequential workflow steps")
            
            # Process each image as a separate step
            loop = asyncio.get_event_loop()

            for idx, file in enumerate(files, 1):
                print(f"\n{'='*60}")
                print(f"PROCESSING STEP {idx}/{len(files)}: {file.filename}")
                print(f"{'='*60}")
                
                # Save image
                image_path = f"input/images/{file.filename}"
                with open(image_path, "wb") as f:
                    file_content = await file.read()
                    f.write(file_content)
                print(f"[STEP {idx}] Saved image to {image_path}, size: {len(file_content)} bytes")
                
                # Read and analyze image with step-specific prompt
                with open(image_path, "rb") as f:
                    image_data = f.read()

                print(f"[STEP {idx}] Calling vision model for analysis...")
                
                # Run vision/network call in thread pool (I/O bound)
#             analysis_prompt = f"""This is Step {idx} of a multi-step workflow process.

# Analyze this image as a SINGLE PROCESS STEP and describe:
# 1. What specific action or process is shown in this step
# 2. What inputs or prerequisites are needed
# 3. What the expected output or result is
# 4. Any important details, settings, or configurations visible
# 5. Any warnings, errors, or special notes

# Provide a clear, detailed description of THIS SPECIFIC STEP ONLY.
# Format your response as a numbered list of observations."""

# Update the analysis prompt in the main function:

                analysis_prompt = f"""This is Step {idx} of a multi-step workflow process.

Analyze this screenshot and extract ALL visible steps, actions, fields, and buttons shown.

For EACH interactive element (field, dropdown, button, etc.), describe:
- The element name/label
- What action to take
- Any specific values or options to select
- The order of operations

Format your response as a detailed numbered list of substeps.
Example:
1. Select Database - Choose 'EMR' or 'GATI' from the dropdown
2. Select Company - Use the dropdown to select the appropriate company
3. Select Date Range - Choose 'From Date' and 'To Date' using calendar picker
Format your response as a numbered list of observations.
Provide a clear, detailed description of THIS SPECIFIC STEP ONLY.
Format your response as a numbered list of observations.
make it as concise as possible."""


                analysis = await loop.run_in_executor(
                    app.state.io_executor,
                    call_ollama_vision,
                    image_data,
                    analysis_prompt,
                )

                print(f"[STEP {idx}] Analysis Result:")
                if analysis:
                    print(analysis)
                    image_analyses.append({
                        'step_number': idx,
                        'filename': file.filename,
                        'image_path': image_path,
                        'analysis': analysis,
                        'status': 'success'
                    })
                else:
                    print(f"[STEP {idx}] No analysis returned")
                    image_analyses.append({
                        'step_number': idx,
                        'filename': file.filename,
                        'image_path': image_path,
                        'analysis': f"Unable to analyze step {idx}",
                        'status': 'error'
                    })

                # Create process tracking image (CPU bound) in process pool
                img_filename = await loop.run_in_executor(
                    app.state.cpu_executor,
                    create_process_image,
                    idx,
                    f"Step {idx}: {file.filename}",
                    f"Analyzed workflow step from {file.filename}",
                    "success" if analysis else "error",
                    {"Image": file.filename, "Analysis Length": f"{len(analysis) if analysis else 0} chars"}
                )

                process_steps.append({
                    "step": idx,
                    "title": f"Step {idx}: {file.filename}",
                    "status": "success" if analysis else "error",
                    "image": f"/static/process_images/{img_filename}"
                })

            print(f"\n[MAIN] Completed analysis of all {len(files)} workflow steps")

        elif description:
            print(f"\n[MAIN] ERROR: Text description not supported for sequential workflow")
            raise HTTPException(status_code=400, detail="Please upload images for sequential workflow analysis")
        else:
            print("[MAIN] ERROR: No input provided")
            raise HTTPException(status_code=400, detail="No images provided")

        # Extract workflow steps from analyses
        print("\n--- Extracting Sequential Workflow Steps ---")
        workflow_steps = extract_steps_from_analysis(image_analyses)
        
        # Create SOP DOCX with images (CPU bound) in process pool
        print("\n--- Creating SOP DOCX with Images ---")
        sop_docx_path = await loop.run_in_executor(
            app.state.cpu_executor,
            create_sop_docx_with_images,
            workflow_steps,
        )
        print(f"SOP DOCX created: {sop_docx_path}")

        print("\n=== PROCESS COMPLETE ===\n")

        return {
            "sop_docx_url": f"/static/{os.path.basename(sop_docx_path)}",
            "process_steps": process_steps,
            "workflow_steps": workflow_steps,
            "total_steps": len(workflow_steps),
            "input_images_count": len(files) if files else 0,
            "message": "Sequential workflow SOP generated successfully with concise summaries",
            "all_success": all(step["status"] == "success" for step in process_steps)
        }


@app.get("/concurrency-status/")
async def concurrency_status():
    """Return the current concurrency usage and configured limits."""
    sem = app.state.semaphore
    # approximate used slots
    used = MAX_CONCURRENT_REQUESTS - sem._value  # note: _value is internal but useful for quick status
    return {
        "max_concurrent_requests": MAX_CONCURRENT_REQUESTS,
        "used_slots": used,
        "available_slots": sem._value,
    }


@app.post("/simulate-parallel/")
async def simulate_parallel(count: int = Form(10), dummy_seconds: int = Form(3)):
    """Simulate 'count' parallel CPU tasks to verify concurrency limit. Each task will run a short CPU-bound work offloaded to the process pool."""
    if count < 1:
        raise HTTPException(status_code=400, detail="count must be >= 1")

    async def _run_dummy(i: int):
        loop = asyncio.get_event_loop()
        def cpu_work(n):
            # simple CPU-bound loop for dummy_seconds seconds
            import time
            end = time.time() + n
            s = 0
            while time.time() < end:
                s += sum(i * i for i in range(1000))
            return s

        return await loop.run_in_executor(app.state.cpu_executor, cpu_work, dummy_seconds)

    tasks = [asyncio.create_task(_run_dummy(i)) for i in range(count)]
    results = await asyncio.gather(*tasks)
    return {"ran": count, "results_sample": results[:3]}

@app.post("/generate/")
async def generate_diagram(
    description: str = Form(None),
    file: UploadFile = File(None),
    output_format: str = Form("png"),
):
    """Simple diagram generation - DEPRECATED for workflow use."""
    raise HTTPException(status_code=400, detail="Please use /generate-process-images/ for workflow generation")

@app.get("/static/{image_name}")
async def get_diagram_image(image_name: str):
    path = f"static/{image_name}"
    if not os.path.exists(path):
        raise HTTPException(404, "Image not found")
    media_type = "image/svg+xml" if image_name.endswith(".svg") else "image/png"
    return FileResponse(path, media_type=media_type)

# Add this to your FastAPI app

@app.post("/regenerate-docx/")
async def regenerate_docx(workflow_steps: List[Dict] = Body(...)):
    """Regenerate DOCX with edited workflow steps."""
    
    print(f"\n[REGENERATE] Regenerating DOCX with {len(workflow_steps)} edited steps")
    
    try:
        # Convert dict to proper format
        formatted_steps = []
        for step in workflow_steps:
            formatted_steps.append({
                "number": step.get("number"),
                "title": step.get("title"),
                "description": step.get("description"),
                "summary": step.get("summary", {}),
                "full_analysis": step.get("full_analysis", ""),
                "image_path": step.get("image_path", ""),
                "filename": step.get("filename", ""),
                "next_step": step.get("next_step")
            })
        
        # Create new DOCX (CPU bound) - offload to process pool
        loop = asyncio.get_event_loop()
        docx_path = await loop.run_in_executor(
            app.state.cpu_executor,
            create_sop_docx_with_images,
            formatted_steps,
        )
        
        return {
            "sop_docx_url": f"/static/{os.path.basename(docx_path)}",
            "message": "DOCX regenerated successfully",
            "total_steps": len(formatted_steps)
        }
        
    except Exception as e:
        print(f"[REGENERATE] Error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))



@app.get("/")
async def root():
    return {"message": "Ollama Sequential Workflow SOP Generator - Upload images as sequential steps"}
