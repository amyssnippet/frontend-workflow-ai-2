import subprocess
import os
import requests
import json
import re
import mimetypes
import base64
import io
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont

import pdfplumber
from docx import Document as DocxDocument
from pptx import Presentation

app = FastAPI(title="Ollama Mermaid Diagram API")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_methods=["*"],
    allow_headers=["*"],
)

if not os.path.exists("static"):
    os.makedirs("static")
app.mount("/static", StaticFiles(directory="static"), name="static")

OLLAMA_API_URL = "https://suggested-accuracy-brokers-gossip.trycloudflare.com/api/chat"
OLLAMA_MODEL = "granite3.3:8b"
VISION_MODEL = "qwen2.5vl:3b"
token = "ac01d6d03fe9feac354add11a4ac67ef51a089594835d56056436f48c9e00ed1"

def check_mmdc_installed():
    try:
        if os.name == "nt":
            mmdc_path = r"C:\Users\amoly\AppData\Roaming\npm\mmdc.cmd"
        else:
            mmdc_path = r"/usr/bin/mmdc"
        subprocess.run([mmdc_path, "--version"], check=True, capture_output=True)
        return True
    except (FileNotFoundError, subprocess.CalledProcessError):
        print("Error: Mermaid CLI (mmdc) not found.")
        return False

def sanitize_mermaid_code(mermaid_code: str) -> str:
    """Clean and normalize Mermaid code to avoid mmdc parse errors."""
    mermaid_code = re.sub(r"^```mermaid\s*\n", "", mermaid_code, flags=re.MULTILINE)
    mermaid_code = re.sub(r"```$", "", mermaid_code, flags=re.MULTILINE)

    lines = mermaid_code.splitlines()
    cleaned = []
    has_graph = False
    node_counter = 0

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("graph "):
            has_graph = True
            if not stripped.endswith(";"):
                stripped += ";"
            cleaned.append(stripped)
            continue

        if any(edge in stripped for edge in ["-->", "-.->", "---"]):
            safe_line = stripped.replace('"', "'")
            safe_line = safe_line.replace("{", "[").replace("}", "]")
            cleaned.append(safe_line)
            continue

        match = re.match(r"^(\w+)\[(.*)\]$", stripped)
        if match:
            node_id, label = match.groups()
            safe_label = label.replace('"', "'")
            safe_label = re.sub(r"[^\w\s\-\.,()']", "", safe_label)
            safe_label = " ".join(safe_label.split())
            cleaned.append(f"{node_id}[{safe_label}]")
            continue

        node_counter += 1
        safe_label = stripped.replace('"', "'")
        safe_label = re.sub(r"[^\w\s\-\.,()']", "", safe_label)
        safe_label = " ".join(safe_label.split())
        cleaned.append(f"N{node_counter}[{safe_label}]")

    if not has_graph:
        cleaned.insert(0, "graph TD;")

    return "\n".join(cleaned) + "\n"

def call_ollama_granite(user_prompt):
    system_message_content = (
        "You are ONLY to output a valid Mermaid flowchart code block. "
        "The output MUST be ONLY the Mermaid code block, enclosed in triple backticks with 'mermaid'. "
        "Rules to follow strictly:\n"
        "1. The diagram MUST start with 'graph TD;' or 'graph LR;'.\n"
        "2. Node text inside [] or {} MUST NOT contain parentheses (), commas, colons, semicolons, or special symbols. "
        "Use simple words only.\n"
        "3. If multiple options are needed, represent them as separate nodes or as edge labels, not inside a single node.\n"
        "4. Only output nodes and edges — no explanations, no comments.\n"
        "If input cannot be converted, output a minimal valid diagram: "
        "``````"
    )

    messages = [
        {"role": "control", "content": "thinking"},
        {"role": "system", "content": system_message_content},
        {"role": "user", "content": user_prompt},
    ]
    payload = {"model": OLLAMA_MODEL, "messages": messages, "stream": False}
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {token}"}

    try:
        response = requests.post(OLLAMA_API_URL, headers=headers, data=json.dumps(payload))
        response.raise_for_status()
        result = response.json()
        print("Ollama response:", result)
        generated_content = result.get("message", {}).get("content", "").strip()

        match = re.search(r"``````", generated_content, re.DOTALL)
        if match:
            return match.group(1).strip()
        else:
            return "graph TD;\nA[No valid Mermaid diagram generated];"
    except Exception as e:
        print(f"Error calling Ollama: {e}")
        return "graph TD;\nA[Error generating diagram];"

def repair_mermaid_with_ollama(broken_code: str) -> str:
    """Ask Ollama to fix invalid Mermaid code."""
    prompt = (
        "The following Mermaid code is invalid. Please fix and return ONLY valid Mermaid code:\n\n"
        f"``````"
    )
    return call_ollama_granite(prompt)

def translate_mermaid_to_image(mermaid_code: str, output_filename: str, output_format="png") -> tuple:
    if not check_mmdc_installed():
        return False, "Mermaid CLI not found"

    temp_mermaid_file = "temp.mmd"
    output_path = f"static/{output_filename}.{output_format}"

    try:
        with open(temp_mermaid_file, "w", encoding="utf-8") as f:
            f.write(mermaid_code)

        if os.name == "nt":
            mmdc_path = r"C:\Users\amoly\AppData\Roaming\npm\mmdc.cmd"
        else:
            mmdc_path = r"/usr/bin/mmdc"

        subprocess.run([mmdc_path, "-i", temp_mermaid_file, "-o", output_path], check=True)
        return True, output_path
    except subprocess.CalledProcessError as e:
        return False, f"mmdc parse error: {e}"
    except Exception as e:
        return False, str(e)
    finally:
        if os.path.exists(temp_mermaid_file):
            os.remove(temp_mermaid_file)

def extract_text_from_file(file_path: str, mime_type: str) -> str:
    """Extract text using multiple methods including vision model."""
    
    def call_vision_model(image_data: bytes) -> str:
        """Call Ollama vision model for text extraction."""
        try:
            if len(image_data) == 0:
                return ""
            
            encoded_image = base64.b64encode(image_data).decode('utf-8')
            print(f"Encoded image size: {len(encoded_image)} chars")
            
            messages = [
                {
                    "role": "user",
                    "content": "Extract all text from this document image. Include headers, body text, lists, tables, and any other visible text. Preserve the structure and formatting. Return only the extracted text without explanations.",
                    "images": [encoded_image]
                }
            ]
            
            payload = {
                "model": VISION_MODEL,
                "messages": messages,
                "stream": False,
                "options": {
                    "temperature": 0.1,
                    "num_ctx": 8192
                }
            }
            
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {token}"
            }
            
            print(f"Calling vision model: {VISION_MODEL}")
            response = requests.post(
                OLLAMA_API_URL, 
                headers=headers, 
                data=json.dumps(payload), 
                timeout=180
            )
            
            print(f"Vision response status: {response.status_code}")
            
            if response.status_code == 200:
                result = response.json()
                extracted = result.get("message", {}).get("content", "").strip()
                print(f"Vision extracted {len(extracted)} characters")
                return extracted
            else:
                print(f"Vision error: {response.text[:500]}")
                return ""
            
        except Exception as e:
            print(f"Vision model error: {e}")
            return ""
    
    def render_docx_to_image(docx_path: str) -> bytes:
        """Render DOCX to image using PIL."""
        try:
            doc = DocxDocument(docx_path)
            
            estimated_lines = len(doc.paragraphs) + sum(len(t.rows) for t in doc.tables)
            img_height = max(1200, estimated_lines * 40 + 100)
            
            img = Image.new('RGB', (1200, img_height), 'white')
            draw = ImageDraw.Draw(img)
            
            try:
                font_normal = ImageFont.truetype("arial.ttf", 24)
            except:
                font_normal = ImageFont.load_default()
            
            y_position = 30
            x_margin = 30
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text = para.text[:150]
                    draw.text((x_margin, y_position), text, fill='black', font=font_normal)
                    y_position += 35
            
            for table in doc.tables:
                y_position += 10
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip()[:20] for cell in row.cells if cell.text.strip()])
                    if row_text:
                        draw.text((x_margin, y_position), row_text, fill='black', font=font_normal)
                        y_position += 30
                y_position += 10
            
            if y_position < img_height:
                img = img.crop((0, 0, 1200, min(y_position + 50, img_height)))
            
            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format='JPEG', quality=70, optimize=True)
            return img_byte_arr.getvalue()
            
        except Exception as e:
            print(f"Error rendering DOCX: {e}")
            return b""
    
    text = ""
    
    try:
        if mime_type == "application/pdf":
            print("Processing PDF...")
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += f"{page_text}\n"
                    else:
                        page_image = page.to_image(resolution=120)
                        img_byte_arr = io.BytesIO()
                        page_image.original.save(img_byte_arr, format='PNG')
                        vision_text = call_vision_model(img_byte_arr.getvalue())
                        if vision_text:
                            text += f"{vision_text}\n"
                    
        elif mime_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]:
            print("Processing DOCX...")
            doc = DocxDocument(file_path)
            
            paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        tables_text.append(row_text)
            
            all_text = paragraphs_text + tables_text
            text = "\n".join(all_text)
            
            if not text.strip():
                rendered_image = render_docx_to_image(file_path)
                if rendered_image:
                    text = call_vision_model(rendered_image)
                
        elif mime_type == "text/plain":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
                
        elif mime_type in ["image/png", "image/jpeg", "image/jpg", "image/bmp", "image/gif"]:
            with open(file_path, "rb") as f:
                text = call_vision_model(f.read())
        else:
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            except:
                text = ""
                
    except Exception as e:
        print(f"Extraction failed: {e}")
        return "Error: Unable to process this document format."
    
    result = text.strip()
    print(f"Final extracted text length: {len(result)}")
    
    if not result:
        return "This document appears to be empty or contains only images/formatting."
    
    return result

@app.post("/generate-docx/")
async def generate_diagram_with_docx(
    description: str = Form(None),
    file: UploadFile = File(None),
    output_format: str = Form("png"),
):
    """Generate a flowchart and create a DOCX document with detailed steps."""
    
    # Step tracking
    steps_log = []
    
    # Step 1: Get input
    prompt_text = description.strip() if description else None
    
    if file is not None:
        steps_log.append("Step 1: File uploaded - extracting text from document")
        input_dir = "input"
        os.makedirs(input_dir, exist_ok=True)
        file_path = os.path.join(input_dir, file.filename)
        
        with open(file_path, "wb") as f:
            f.write(await file.read())
        
        mime_type, _ = mimetypes.guess_type(file_path)
        if not mime_type:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        
        extracted_text = extract_text_from_file(file_path, mime_type)
        if not extracted_text:
            raise HTTPException(status_code=422, detail="File is empty or unsupported")
        
        prompt_text = extracted_text
        steps_log.append(f"Step 1 Complete: Extracted {len(extracted_text)} characters from {file.filename}")
    else:
        steps_log.append("Step 1: Using provided text description")

    if not prompt_text:
        raise HTTPException(status_code=400, detail="No input provided")

    # Step 2: Generate Mermaid code
    steps_log.append(f"Step 2: Generating Mermaid flowchart code using {OLLAMA_MODEL}")
    mermaid_code = call_ollama_granite(prompt_text)
    mermaid_code_original = mermaid_code
    mermaid_code = sanitize_mermaid_code(mermaid_code)
    steps_log.append("Step 2 Complete: Mermaid code generated and sanitized")

    # Step 3: Render to image
    steps_log.append("Step 3: Rendering Mermaid code to image")
    success, result = translate_mermaid_to_image(mermaid_code, "generated_flowchart", output_format)

    # Step 4: Repair if needed
    if not success:
        steps_log.append("Step 3 Failed: Initial render failed, attempting repair")
        steps_log.append(f"Step 4: Repairing Mermaid code using {OLLAMA_MODEL}")
        repaired_code = repair_mermaid_with_ollama(mermaid_code)
        repaired_code = sanitize_mermaid_code(repaired_code)
        success, result = translate_mermaid_to_image(repaired_code, "generated_flowchart", output_format)

        if success:
            mermaid_code = repaired_code
            steps_log.append("Step 4 Complete: Repair successful, diagram generated")
        else:
            mermaid_code = "graph TD;\nA[Diagram generation failed];"
            result = f"static/generated_flowchart.{output_format}"
            steps_log.append("Step 4 Failed: Unable to generate valid diagram")
    else:
        steps_log.append("Step 3 Complete: Diagram successfully rendered")

    # Step 5: Create DOCX
    steps_log.append("Step 5: Creating comprehensive DOCX document")
    
    doc = DocxDocument()
    
    # Title
    title = doc.add_heading('Flowchart Generation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Processing Steps Section
    doc.add_heading('Processing Steps', level=1)
    for step in steps_log:
        p = doc.add_paragraph(step, style='List Bullet')
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph()  # Spacing
    
    # Input Content Section
    doc.add_heading('Input Content', level=1)
    input_para = doc.add_paragraph()
    input_para.add_run('Source: ').bold = True
    input_para.add_run(f"{file.filename if file else 'Text description'}\n")
    input_para.add_run('Length: ').bold = True
    input_para.add_run(f"{len(prompt_text)} characters\n\n")
    
    # Truncate very long text
    display_text = prompt_text if len(prompt_text) <= 2000 else prompt_text[:2000] + "\n\n[... text truncated for brevity ...]"
    doc.add_paragraph(display_text)
    
    doc.add_page_break()
    
    # Generated Flowchart Section
    doc.add_heading('Generated Flowchart Diagram', level=1)
    
    # Add the flowchart image if it exists
    if os.path.exists(result) and os.path.getsize(result) > 0:
        try:
            doc.add_picture(result, width=Inches(6.0))
            steps_log.append("Step 5: Flowchart image added to document")
        except Exception as e:
            doc.add_paragraph(f"Error adding image: {str(e)}")
            print(f"Error adding image to DOCX: {e}")
    else:
        doc.add_paragraph("⚠ Flowchart image could not be generated or is empty")
    
    doc.add_paragraph()  # Spacing
    
    # Mermaid Code Section
    doc.add_heading('Mermaid Source Code', level=1)
    doc.add_paragraph('This code can be used in any Mermaid-compatible viewer:')
    
    # Add code in a styled paragraph
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(mermaid_code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    code_run.font.color.rgb = RGBColor(0, 0, 128)
    
    # Metadata Section
    doc.add_page_break()
    doc.add_heading('Generation Metadata', level=1)
    
    metadata = [
        f"Model Used: {OLLAMA_MODEL}",
        f"Vision Model: {VISION_MODEL}",
        f"Output Format: {output_format.upper()}",
        f"Mermaid Code Length: {len(mermaid_code)} characters",
        f"Repair Required: {'Yes' if not success else 'No'}",
        f"Total Processing Steps: {len(steps_log)}"
    ]
    
    for item in metadata:
        doc.add_paragraph(item, style='List Bullet')
    
    # Save DOCX
    docx_filename = "flowchart_generation_report.docx"
    docx_path = f"static/{docx_filename}"
    doc.save(docx_path)
    
    steps_log.append("Step 5 Complete: DOCX document created successfully")
    
    print(f"DOCX saved to: {docx_path}")
    print(f"Image path: {result}")
    print(f"Image exists: {os.path.exists(result)}")
    print(f"Image size: {os.path.getsize(result) if os.path.exists(result) else 0} bytes")

    return {
        "mermaid": mermaid_code,
        "image_url": f"/static/{os.path.basename(result)}",
        "docx_url": f"/static/{docx_filename}",
        "message": "Flowchart and detailed report generated successfully",
        "steps": steps_log,
        "image_exists": os.path.exists(result),
        "image_size": os.path.getsize(result) if os.path.exists(result) else 0
    }

@app.post("/generate/")
async def generate_diagram(
    description: str = Form(None),
    file: UploadFile = File(None),
    output_format: str = Form("png"),
):
    prompt_text = description.strip() if description else None

    if file is not None:
        input_dir = "input"
        os.makedirs(input_dir, exist_ok=True)
        file_path = os.path.join(input_dir, file.filename)
        with open(file_path, "wb") as f:
            f.write(await file.read())
        mime_type, _ = mimetypes.guess_type(file_path)
        if not mime_type:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        extracted_text = extract_text_from_file(file_path, mime_type)
        if not extracted_text:
            raise HTTPException(status_code=422, detail="File is empty or unsupported")
        prompt_text = extracted_text

    if not prompt_text:
        raise HTTPException(status_code=400, detail="No input provided")

    mermaid_code = call_ollama_granite(prompt_text)
    mermaid_code = sanitize_mermaid_code(mermaid_code)

    success, result = translate_mermaid_to_image(mermaid_code, "generated_flowchart", output_format)

    if not success:
        repaired_code = repair_mermaid_with_ollama(mermaid_code)
        repaired_code = sanitize_mermaid_code(repaired_code)
        success, result = translate_mermaid_to_image(repaired_code, "generated_flowchart", output_format)

        if success:
            mermaid_code = repaired_code
        else:
            mermaid_code = "graph TD;\nA[Diagram generation failed];"
            result = f"static/generated_flowchart.{output_format}"
            with open(result, "wb") as f:
                f.write(b"")

    return {"mermaid": mermaid_code, "image_url": f"/static/{os.path.basename(result)}"}

@app.get("/static/{image_name}")
async def get_diagram_image(image_name: str):
    path = f"static/{image_name}"
    if not os.path.exists(path):
        raise HTTPException(404, "Image not found")
    media_type = "image/svg+xml" if image_name.endswith(".svg") else "image/png"
    return FileResponse(path, media_type=media_type)

@app.get("/")
async def root():
    return {"message": "Ollama Mermaid FastAPI service running with enhanced DOCX generation"}
